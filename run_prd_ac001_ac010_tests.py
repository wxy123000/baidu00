import ast
import json
import re
import zipfile
from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader
from reportlab.pdfgen import canvas
from streamlit.testing.v1 import AppTest
from xml.etree import ElementTree


ROOT = Path(__file__).resolve().parent
APP = ROOT / "app.py"
OUT = ROOT / ".mvp_test_runtime" / "ac001_010"
OUT.mkdir(parents=True, exist_ok=True)
SOURCE = APP.read_text(encoding="utf-8")


def load_app_functions():
    tree = ast.parse(SOURCE)
    wanted = {
        "extract_file",
        "split_long_text",
        "question_terms",
        "select_long_document_context",
    }
    nodes = [
        node for node in tree.body
        if isinstance(node, ast.FunctionDef) and node.name in wanted
    ]
    namespace = {
        "BytesIO": BytesIO,
        "PdfReader": PdfReader,
        "Document": Document,
        "zipfile": zipfile,
        "ElementTree": ElementTree,
        "re": re,
    }
    exec(compile(ast.Module(body=nodes, type_ignores=[]), str(APP), "exec"), namespace)
    return namespace


class Upload:
    def __init__(self, name, data):
        self.name = name
        self._data = data
        self.size = len(data)

    def getvalue(self):
        return self._data


def make_pdf():
    buffer = BytesIO()
    c = canvas.Canvas(buffer)
    c.drawString(72, 760, "PDF_TOKEN_2026 OWNER_ZHANG BUDGET_50000")
    c.save()
    return buffer.getvalue()


def make_docx():
    buffer = BytesIO()
    doc = Document()
    doc.add_paragraph("WORD_TOKEN_2026 OWNER_LI PERIOD_4_WEEKS")
    doc.save(buffer)
    return buffer.getvalue()


def make_pptx():
    buffer = BytesIO()
    slide_template = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:sld xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
 xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main">
 <p:cSld><p:spTree><p:sp><p:txBody><a:p><a:r><a:t>{text}</a:t></a:r></a:p>
 </p:txBody></p:sp></p:spTree></p:cSld></p:sld>"""
    with zipfile.ZipFile(buffer, "w") as package:
        package.writestr("ppt/slides/slide1.xml", slide_template.format(text="PPT_TOKEN_2026 PROJECT_ALPHA"))
        package.writestr("ppt/slides/slide2.xml", slide_template.format(text="OWNER_WANG BUDGET_80000"))
        package.writestr("ppt/slides/slide3.xml", slide_template.format(text="PERIOD_6_WEEKS RELEASE_0915"))
    return buffer.getvalue()


def result(case_id, title, passed, evidence):
    return {"case_id": case_id, "title": title, "passed": bool(passed), "evidence": evidence}


ns = load_app_functions()
extract_file = ns["extract_file"]
select_context = ns["select_long_document_context"]
results = []

# AC-001 and AC-002: real Streamlit page flow.
at = AppTest.from_file(str(APP), default_timeout=30)
at.run()
home_ok = len(at.exception) == 0 and at.session_state["page"] == "首页"
home_ok = home_ok and any(x.label == "快速输入" for x in at.text_area)
home_ok = home_ok and at.sidebar.radio[0].value == "首页"
results.append(result("AC-001", "启动与首页", home_ok,
                      f"exceptions={len(at.exception)}, page={at.session_state['page']}, nav={at.sidebar.radio[0].value}"))

quick_input = next(x for x in at.text_area if x.label == "快速输入")
quick_input.set_value("AC002手动文本：项目负责人张明，测试周期4周。")
next(x for x in at.button if x.label == "📋 使用文本").click().run()
ac002_ok = len(at.exception) == 0 and at.session_state["page"] == "选择功能"
ac002_ok = ac002_ok and "AC002手动文本" in at.session_state["source_text"]
results.append(result("AC-002", "手动输入文本", ac002_ok,
                      f"page={at.session_state['page']}, source_len={len(at.session_state['source_text'])}"))

# Real extraction samples for supported formats.
samples = {
    "pdf": Upload("sample.pdf", make_pdf()),
    "docx": Upload("sample.docx", make_docx()),
    "pptx": Upload("sample.pptx", make_pptx()),
    "txt": Upload("sample.txt", b"TXT_TOKEN_2026 USERS_20"),
    "md": Upload("sample.md", b"# MD_TOKEN_2026\nSTATUS_OK"),
}
parsed = {name: extract_file(upload) for name, upload in samples.items()}
accepted_types = re.search(r'type=\[([^\]]+)\]', SOURCE).group(1)
ac003_ok = all(parsed.values()) and all(f'"{ext}"' in accepted_types for ext in samples)
results.append(result("AC-003", "支持规定文件格式", ac003_ok,
                      {name: len(text) for name, text in parsed.items()}))

max_bytes_match = re.search(r"MAX_UPLOAD_BYTES\s*=\s*20\s*\*\s*1024\s*\*\s*1024", SOURCE)
oversize_guard = "> MAX_UPLOAD_BYTES" in SOURCE and "以下文件超过 20MB" in SOURCE
results.append(result("AC-004", "单文件不超过20MB", bool(max_bytes_match and oversize_guard),
                      "limit=20*1024*1024; oversize guard and message present"))

max_files_match = re.search(r"MAX_UPLOAD_FILES\s*=\s*3", SOURCE)
count_guard = "len(uploads or []) > MAX_UPLOAD_FILES" in SOURCE
results.append(result("AC-005", "单次上传1至3个文件", bool(max_files_match and count_guard),
                      "limit=3; fourth-file guard present"))

status_ok = "✓ 待解析" in SOURCE and "✕ 超过 20MB" in SOURCE
name_size_ok = "file.name" in SOURCE and "size / 1024 / 1024" in SOURCE
results.append(result("AC-006", "显示文件名称和状态", status_ok and name_size_ok,
                      "name, size, pending/oversize status rendering present"))

# AC-007: unsupported, damaged and empty files.
unsupported_blocked = '"xlsx"' not in accepted_types
damaged_failed = False
try:
    extract_file(Upload("damaged.pdf", b"this is not a pdf"))
except Exception:
    damaged_failed = True
empty_not_extracted = extract_file(Upload("empty.txt", b"")) == ""
error_messages = (
    "文件解析失败" in SOURCE
    and "未提取到可读文字" in SOURCE
    and "请至少上传一个可读取的文件" in SOURCE
)
results.append(result("AC-007", "异常文件提示", unsupported_blocked and damaged_failed and empty_not_extracted and error_messages,
                      f"unsupported={unsupported_blocked}, damaged_error={damaged_failed}, empty={empty_not_extracted}, messages={error_messages}"))

# AC-008 and AC-009: assert exact body tokens were extracted.
ac008_ok = (
    "PDF_TOKEN_2026" in parsed["pdf"]
    and "WORD_TOKEN_2026" in parsed["docx"]
    and "TXT_TOKEN_2026" in parsed["txt"]
)
results.append(result("AC-008", "解析PDF、Word和TXT正文", ac008_ok,
                      "PDF_TOKEN_2026, WORD_TOKEN_2026, TXT_TOKEN_2026 extracted"))

ac009_ok = all(token in parsed["pptx"] for token in (
    "PPT_TOKEN_2026", "OWNER_WANG", "PERIOD_6_WEEKS", "RELEASE_0915"
)) and all(f"【第 {i} 页】" in parsed["pptx"] for i in (1, 2, 3))
results.append(result("AC-009", "解析PowerPoint正文", ac009_ok,
                      "three slide labels and all unique tokens extracted"))

# AC-010: use the supplied long DOCX and ask three independent questions.
long_doc = Document(ROOT / "AC-010长文档分段处理测试样本.docx")
long_text = "\n".join(p.text for p in long_doc.paragraphs)
queries = [
    ("文档开头的项目代号和启动负责人是什么？", "START-2026"),
    ("文档中间的核心验收编号和中期预算是什么？", "MIDDLE-8888"),
    ("文档结尾的最终发布编号和发布日期是什么？", "END-9999"),
]
retrieval = []
for question, marker in queries:
    context = select_context(long_text, question, 30000)
    retrieval.append({"marker": marker, "found": marker in context, "context_len": len(context)})
results.append(result("AC-010", "长文档分段检索", all(x["found"] for x in retrieval),
                      {"document_len": len(long_text), "retrieval": retrieval}))

summary = {
    "total": len(results),
    "passed": sum(item["passed"] for item in results),
    "failed": sum(not item["passed"] for item in results),
    "results": results,
}
(ROOT / "prd_ac001_ac010_results.json").write_text(
    json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8"
)
print(json.dumps(summary, ensure_ascii=False, indent=2))
