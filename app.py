from datetime import datetime
from io import BytesIO
import json
import os
import re
import sqlite3
import time
import uuid
import zipfile
from xml.etree import ElementTree
from xml.sax.saxutils import escape

import streamlit as st
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from openai import OpenAI
from pypdf import PdfReader
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


st.set_page_config(
    page_title="百度办公 AI 助手",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded",
)


FEATURES = {
    "会议纪要": {
        "icon": "📝",
        "desc": "自动提炼会议摘要、关键结论和待办事项。",
        "tags": "会议记录　讨论文本",
    },
    "文档问答": {
        "icon": "💬",
        "desc": "基于上传文档回答问题，并定位原文依据。",
        "tags": "PDF　Word　长文档",
    },
    "任务拆解": {
        "icon": "🧩",
        "desc": "将项目目标拆为任务，给出负责人、优先级和时间建议。",
        "tags": "项目计划　任务安排",
    },
    "周报生成": {
        "icon": "📊",
        "desc": "汇总本周工作内容，自动生成结构清晰的周报。",
        "tags": "工作记录　周报模板",
    },
}


def init_state():
    defaults = {
        "page": "首页",
        "source_text": "",
        "source_names": [],
        "selected_feature": "会议纪要",
        "workflow_feature_active": False,
        "question": "",
        "result": "",
        "history": [],
        "saved_files": [],
        "ai_provider": "本地演示",
        "ai_error": "",
        "recommended_feature": "",
        "recommendation_reason": "",
        "chat_messages": [
            {"role": "assistant", "content": "你好！我可以直接回答办公问题，也可以结合已上传资料连续对话。"}
        ],
        "automation_tasks": [],
        "automation_provider": "待生成",
        "session_id": str(uuid.uuid4()),
        "user_id": str(uuid.uuid4()),
        "visit_logged": False,
        "last_generation_ms": 0,
        "result_confirmed": False,
        "persistent_content_loaded": False,
        "confirmation_notice": False,
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


DB_PATH = os.path.join(os.path.dirname(__file__), "mvp_metrics.db")
MAX_UPLOAD_FILES = 3
MAX_UPLOAD_BYTES = 20 * 1024 * 1024


def init_database():
    with sqlite3.connect(DB_PATH) as connection:
        connection.execute(
            """CREATE TABLE IF NOT EXISTS events (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                created_at TEXT NOT NULL,
                user_id TEXT NOT NULL,
                session_id TEXT NOT NULL,
                event_name TEXT NOT NULL,
                feature TEXT,
                success INTEGER,
                duration_ms INTEGER,
                detail TEXT
            )"""
        )
        connection.execute(
            """CREATE TABLE IF NOT EXISTS feedback (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                created_at TEXT NOT NULL,
                user_id TEXT NOT NULL,
                feature TEXT,
                rating INTEGER NOT NULL,
                helpful INTEGER NOT NULL,
                reuse_intent INTEGER NOT NULL,
                comment TEXT
            )"""
        )
        connection.execute(
            """CREATE TABLE IF NOT EXISTS history_records (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                feature TEXT NOT NULL,
                title TEXT NOT NULL,
                source TEXT NOT NULL,
                created_at TEXT NOT NULL,
                result TEXT NOT NULL
            )"""
        )
        connection.execute(
            """CREATE TABLE IF NOT EXISTS saved_result_files (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL,
                data BLOB NOT NULL,
                mime TEXT NOT NULL,
                created_at TEXT NOT NULL
            )"""
        )


def load_persistent_content():
    if st.session_state.persistent_content_loaded:
        return
    with sqlite3.connect(DB_PATH) as connection:
        history_rows = connection.execute(
            "SELECT id, feature, title, source, created_at, result FROM history_records ORDER BY id DESC LIMIT 100"
        ).fetchall()
        file_rows = connection.execute(
            "SELECT id, name, data, mime, created_at FROM saved_result_files ORDER BY id ASC LIMIT 100"
        ).fetchall()
    st.session_state.history = [
        {"id": row[0], "feature": row[1], "title": row[2], "source": row[3], "time": row[4], "result": row[5]}
        for row in history_rows
    ]
    st.session_state.saved_files = [
        {"id": row[0], "name": row[1], "data": bytes(row[2]), "mime": row[3], "time": row[4]}
        for row in file_rows
    ]
    st.session_state.persistent_content_loaded = True


def delete_history_record(record_id):
    if record_id is None:
        return
    with sqlite3.connect(DB_PATH) as connection:
        connection.execute("DELETE FROM history_records WHERE id = ?", (record_id,))


def clear_history_records():
    with sqlite3.connect(DB_PATH) as connection:
        connection.execute("DELETE FROM history_records")


def persist_saved_file(item):
    with sqlite3.connect(DB_PATH) as connection:
        cursor = connection.execute(
            "INSERT INTO saved_result_files (name, data, mime, created_at) VALUES (?, ?, ?, ?)",
            (item["name"], sqlite3.Binary(item["data"]), item["mime"], item["time"]),
        )
    item["id"] = cursor.lastrowid
    return item


def log_event(event_name, feature="", success=None, duration_ms=None, detail=""):
    """Store anonymous product events without saving uploaded document content."""
    with sqlite3.connect(DB_PATH) as connection:
        connection.execute(
            """INSERT INTO events
               (created_at, user_id, session_id, event_name, feature, success, duration_ms, detail)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
            (
                datetime.now().isoformat(timespec="seconds"),
                st.session_state.user_id,
                st.session_state.session_id,
                event_name,
                feature or None,
                None if success is None else int(bool(success)),
                duration_ms,
                detail[:500] or None,
            ),
        )


def save_feedback(rating, helpful, reuse_intent, comment):
    with sqlite3.connect(DB_PATH) as connection:
        connection.execute(
            """INSERT INTO feedback
               (created_at, user_id, feature, rating, helpful, reuse_intent, comment)
               VALUES (?, ?, ?, ?, ?, ?, ?)""",
            (
                datetime.now().isoformat(timespec="seconds"),
                st.session_state.user_id,
                st.session_state.selected_feature,
                rating,
                int(helpful),
                int(reuse_intent),
                comment.strip()[:1000],
            ),
        )
    log_event("feedback_submitted", st.session_state.selected_feature, True)


def go(page):
    st.session_state.page = page


def handle_nav_change():
    """在侧栏选项变化时一次性完成导航，避免页面主体按钮事件被重跑吞掉。"""
    chosen_nav = st.session_state.get("nav_radio_v2")
    if not chosen_nav:
        return
    if chosen_nav in FEATURES:
        st.session_state.selected_feature = chosen_nav
        st.session_state.workflow_feature_active = True
        go("选择功能" if st.session_state.source_text else "上传资料")
    else:
        if chosen_nav == "上传资料":
            st.session_state.workflow_feature_active = False
        go(chosen_nav)


def update_automation_task_status(index):
    """Apply a checkbox change before the page recalculates overall progress."""
    tasks = st.session_state.get("automation_tasks", [])
    if index >= len(tasks):
        return
    new_done = bool(st.session_state.get(f"automation_done_{index}", False))
    old_done = bool(tasks[index].get("done"))
    tasks[index]["done"] = new_done
    if new_done != old_done:
        log_event(
            "automation_status_changed",
            "任务自动化",
            True,
            detail=f"index={index};done={int(new_done)}",
        )


def clear_automation_checkbox_state():
    for key in list(st.session_state):
        if str(key).startswith("automation_done_"):
            del st.session_state[key]


def extract_file(uploaded_file):
    suffix = uploaded_file.name.rsplit(".", 1)[-1].lower()
    data = uploaded_file.getvalue()
    if suffix in {"txt", "md"}:
        for encoding in ("utf-8", "gb18030", "utf-16"):
            try:
                return data.decode(encoding)
            except UnicodeDecodeError:
                continue
        return data.decode("utf-8", errors="ignore")
    if suffix == "pdf":
        reader = PdfReader(BytesIO(data))
        pages = []
        for page_number, page in enumerate(reader.pages, 1):
            page_text = (page.extract_text() or "").strip()
            if page_text:
                pages.append(f"【第 {page_number} 页】\n{page_text}")
        return "\n\n".join(pages)
    if suffix == "docx":
        document = Document(BytesIO(data))
        paragraphs = []
        for paragraph_number, paragraph in enumerate(document.paragraphs, 1):
            paragraph_text = paragraph.text.strip()
            if paragraph_text:
                paragraphs.append(f"【段落 {paragraph_number}】\n{paragraph_text}")
        return "\n\n".join(paragraphs)
    if suffix == "pptx":
        # PPTX is an OOXML zip package. Extract slide text without adding a
        # heavyweight presentation dependency to the MVP runtime.
        with zipfile.ZipFile(BytesIO(data)) as package:
            slide_names = sorted(
                (
                    name for name in package.namelist()
                    if name.startswith("ppt/slides/slide") and name.endswith(".xml")
                ),
                key=lambda name: int(name.rsplit("slide", 1)[1].split(".xml", 1)[0]),
            )
            slides = []
            for slide_number, name in enumerate(slide_names, 1):
                root = ElementTree.fromstring(package.read(name))
                texts = [node.text for node in root.iter() if node.tag.endswith("}t") and node.text]
                if texts:
                    slides.append(f"【第 {slide_number} 页】\n" + "\n".join(texts))
            return "\n\n".join(slides)
    return ""


def source_excerpt(text, limit=240):
    clean = " ".join(text.split())
    if not clean:
        return "当前未检测到可用正文，请补充资料后重新生成。"
    return clean[:limit] + ("……" if len(clean) > limit else "")


def split_long_text(text, chunk_size=4000, overlap=400):
    """Split long text into overlapping chunks without losing paragraph context."""
    if not text:
        return []
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        # Prefer a paragraph boundary near the nominal chunk end.
        if end < len(text):
            boundary = text.rfind("\n", start + chunk_size // 2, end)
            if boundary > start:
                end = boundary
        chunks.append(text[start:end])
        if end >= len(text):
            break
        start = max(end - overlap, start + 1)
    return chunks


def question_terms(question):
    """Build lightweight Chinese/ASCII terms for deterministic local retrieval."""
    cleaned = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff]+", "", question or "")
    stop_phrases = ("请问", "根据", "文档", "资料", "回答", "分别", "什么", "是多少", "如何", "是否")
    for phrase in stop_phrases:
        cleaned = cleaned.replace(phrase, "")
    terms = set(re.findall(r"[A-Za-z0-9_-]{2,}", cleaned))
    chinese = "".join(re.findall(r"[\u4e00-\u9fff]", cleaned))
    for size in (2, 3, 4, 5, 6):
        terms.update(chinese[i:i + size] for i in range(max(0, len(chinese) - size + 1)))
    return {term for term in terms if len(term) >= 2}


def select_long_document_context(text, question="", limit=30000):
    """Keep long documents searchable instead of dropping their middle section."""
    if len(text) <= limit:
        return text
    chunks = split_long_text(text)
    terms = question_terms(question)
    scored = []
    for index, chunk in enumerate(chunks):
        lowered = chunk.lower()
        score = sum((len(term) ** 2) * lowered.count(term.lower()) for term in terms)
        scored.append((score, index, chunk))

    # Preserve document framing and add the most question-relevant segments.
    selected = {0, len(chunks) - 1}
    for score, index, _ in sorted(scored, key=lambda item: (-item[0], item[1])):
        if score <= 0 and len(selected) >= 4:
            break
        selected.add(index)
        if len(selected) >= 7:
            break

    parts = []
    used = 0
    for index in sorted(selected):
        labelled = f"【文档片段 {index + 1}/{len(chunks)}】\n{chunks[index].strip()}"
        if used + len(labelled) > limit:
            remaining = limit - used
            if remaining > 200:
                parts.append(labelled[:remaining])
            break
        parts.append(labelled)
        used += len(labelled) + 2
    return "\n\n".join(parts)


def retrieve_relevant_passages(text, question, max_passages=3):
    """Return relevant passages with source file and page/segment labels."""
    context = select_long_document_context(text, question, 30000)
    terms = question_terms(question)
    candidates = []
    current_source = "当前资料"
    current_location = ""
    order = 0
    for raw_line in context.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        label = re.fullmatch(r"【([^】]+)】", line)
        if label:
            value = label.group(1).strip()
            if value == "粘贴文本" or re.search(r"\.(pdf|docx|pptx|txt|md)$", value, re.IGNORECASE):
                current_source = value
                current_location = ""
            elif value.startswith(("第 ", "段落 ", "文档片段 ")):
                current_location = value
            continue
        for raw in re.split(r"[。！？]+", line):
            passage = raw.strip()
            if not passage:
                continue
            lowered = passage.lower()
            score = sum((len(term) ** 2) * lowered.count(term.lower()) for term in terms)
            if score > 0:
                candidates.append((score, order, current_source, current_location, passage[:600]))
            order += 1
    ranked = sorted(candidates, key=lambda item: (-item[0], item[1]))
    selected = []
    for _, _, source, location, passage in ranked:
        label = f"来源：{source}"
        if location:
            label += f"｜{location}"
        formatted = f"【{label}】{passage}"
        if formatted not in selected:
            selected.append(formatted)
        if len(selected) >= max_passages:
            break
    if selected:
        return selected
    fallback = source_excerpt(context, 500)
    return [f"【来源：{current_source}｜相关位置未定位】{fallback}"]


def question_has_document_answer(text, question):
    """Conservatively detect whether common requested fields exist in the document."""
    source = text or ""
    asked = question or ""
    field_signals = {
        "负责人": ("负责人", "负责", "责任人", "owner"),
        "截止时间": ("截止", "完成时间", "完成日期", "deadline"),
        "预算": ("预算", "金额", "费用", "经费"),
        "测试周期": ("测试周期", "周期", "测试时间"),
        "测试人数": ("测试人数", "测试用户", "测试人员", "用户数量"),
        "发布日期": ("发布日期", "发布时间", "上线时间", "发布日"),
    }
    requested = [signals for field, signals in field_signals.items() if field in asked]
    if requested and any(not any(signal.lower() in source.lower() for signal in signals) for signals in requested):
        return False
    terms = question_terms(asked)
    if not terms:
        return bool(source.strip())
    return any(term.lower() in source.lower() for term in terms)


def redact_sensitive_text(text):
    """Mask common personal and credential data before model use or display."""
    if not text:
        return text

    redacted = str(text)
    # Chinese resident identity number: keep only the first 3 and last 4 chars.
    redacted = re.sub(
        r"(?<!\d)(\d{3})\d{11}(\d{3}[\dXx])(?!\d)",
        lambda match: f"{match.group(1)}***********{match.group(2)}",
        redacted,
    )
    # Mainland mobile number.
    redacted = re.sub(
        r"(?<!\d)(1[3-9]\d)\d{4}(\d{4})(?!\d)",
        lambda match: f"{match.group(1)}****{match.group(2)}",
        redacted,
    )
    # Typical bank-card number. Identity numbers have already been masked above.
    redacted = re.sub(
        r"(?<!\d)(\d{4})\d{8,11}(\d{4})(?!\d)",
        lambda match: f"{match.group(1)} **** **** {match.group(2)}",
        redacted,
    )
    # Password values explicitly attached to a password label.
    redacted = re.sub(
        r"(支付密码|登录密码|邮箱密码|账户密码|密码)\s*[:：为是]?\s*"
        r"([A-Za-z0-9!@#$%^&*._-]{4,})",
        lambda match: f"{match.group(1)}：******",
        redacted,
        flags=re.IGNORECASE,
    )
    # Detailed addresses are hidden only when supplied as a labelled value.
    redacted = re.sub(
        r"(家庭地址|详细地址|住址)\s*(?:[:：]|为|是)\s*([^\n，。；;（）()]{4,})",
        lambda match: f"{match.group(1)}：[已隐藏]",
        redacted,
    )
    return redacted


def normalize_generated_text(text):
    """Keep AI output readable by removing accidental large vertical gaps."""
    if not text:
        return text
    normalized = str(text).replace("\r\n", "\n").replace("\r", "\n")
    # Lines containing only spaces should count as blank lines too.
    normalized = re.sub(r"\n[ \t]+\n", "\n\n", normalized)
    # Preserve normal paragraph separation, but collapse 3+ line breaks.
    return re.sub(r"\n{3,}", "\n\n", normalized.strip())


def detect_safety_risk(text, question=""):
    """Return a high-risk request category, or an empty string when allowed."""
    content = f"{text}\n{question}".lower()
    rules = [
        (
            "unauthorized_access",
            [r"获取.{0,12}(邮箱|账户|登录)?密码", r"绕过.{0,12}(权限|访问控制)", r"未经授权.{0,12}(登录|访问)"],
        ),
        (
            "destructive_delete",
            [r"删除(数据库中?)?(的)?全部", r"清空.{0,8}(数据库|记录)", r"不(保留|做).{0,6}备份"],
        ),
        (
            "confidential_exfiltration",
            [r"(机密|保密).{0,24}(发送|提供|导出).{0,12}(外部|外人|第三方)", r"发送给外部人员"],
        ),
        (
            "discrimination",
            [r"排除女性", r"不招聘女性", r"按照性别.{0,8}(排除|拒绝)"],
        ),
        (
            "unverified_allegation",
            [r"未经(核实|证实).{0,16}(违规|违法|传闻)", r"(违规|违法)传闻.{0,10}(确定|事实)"],
        ),
        (
            "conceal_breach",
            [r"(隐藏|隐瞒).{0,12}数据泄露", r"数据泄露.{0,16}(不要|不).{0,8}(记录|上报|通知)"],
        ),
    ]
    for category, patterns in rules:
        if any(re.search(pattern, content, flags=re.IGNORECASE) for pattern in patterns):
            return category
    return ""


def build_safety_response(category):
    responses = {
        "unauthorized_access": (
            "无法执行该请求。获取他人密码、绕过权限或未经授权登录账户会侵犯隐私并带来安全和法律风险。\n\n"
            "合规做法：请联系账户本人或系统管理员，通过授权、共享或权限申请流程获取所需文件。"
        ),
        "destructive_delete": (
            "无法制定“删除全部数据且不备份”的直接执行计划。该操作可能造成不可恢复的数据损失。\n\n"
            "执行任何删除前必须完成：书面审批、目标范围确认、可用备份、恢复演练和操作审计；优先采用可撤销的软删除。"
        ),
        "confidential_exfiltration": (
            "无法协助向外部人员发送公司机密资料或规避访问控制。\n\n"
            "如确有业务需要，请先取得资料所有者和安全负责人的书面授权，并通过公司批准的加密共享渠道处理。"
        ),
        "discrimination": (
            "该招聘提议不能实施。因性别或可能休产假而排除女性候选人属于歧视性做法，并可能违反劳动与平等就业要求。\n\n"
            "应按照岗位能力、经验和统一评价标准开展招聘，并由人力资源或合规人员复核招聘规则。"
        ),
        "unverified_allegation": (
            "不能把未经核实的违规传闻写成确定事实。这样可能损害当事人声誉并带来合规风险。\n\n"
            "建议仅记录为“待核实风险线索”，限制传播范围，并交由有权限的负责人按照正式调查流程处理。"
        ),
        "conceal_breach": (
            "不能隐瞒数据泄露事件或省略记录与上报。\n\n"
            "应立即停止相关操作、保留日志和证据、控制影响范围，并通知项目负责人及安全负责人；完成影响评估、修复和复盘后再恢复相关处理。"
        ),
    }
    return responses[category]


def prepare_safe_input(text, question=""):
    """Apply policy checks and return sanitized content plus an optional block."""
    category = detect_safety_risk(text, question)
    if category:
        return "", "", build_safety_response(category)
    return redact_sensitive_text(text), redact_sensitive_text(question), ""


def generate_local_result(feature, text, question=""):
    safe_text, safe_question, blocked_response = prepare_safe_input(text, question)
    if blocked_response:
        return blocked_response
    excerpt = source_excerpt(safe_text)
    if feature == "会议纪要":
        return f"""会议纪要

一、会议摘要
本次会议围绕上传资料中的工作安排与项目推进情况展开讨论。资料要点：{excerpt}

二、关键结论
1. 明确当前阶段的核心目标与交付范围。
2. 按优先级推进重点事项，并及时同步风险。
3. 重要决策需由相关负责人确认后执行。

三、待办事项
1. 根据会议结论推进后续事项｜负责人：待确认｜优先级：待确认｜截止时间：待确认

四、负责人建议
资料未明确具体负责人，标注为“待确认”，由项目组会后确认。

五、风险与备注
资料中未明确的负责人、优先级和截止时间均未作确定性补充，正式使用前请人工确认。"""
    if feature == "文档问答":
        asked = safe_question.strip() or "这份资料的主要内容是什么？"
        if not question_has_document_answer(safe_text, asked):
            return f"""问题
{asked}

直接答案
无法确定。上传资料中未提及该问题所需的信息。

关键依据
- 已检索当前文档，但未找到能够支持该答案的相关字段或原文。

说明：系统不会使用文档之外的信息补充答案。"""
        passages = retrieve_relevant_passages(safe_text, asked)
        direct_answer = passages[0]
        evidence = "\n".join(f"- {passage}" for passage in passages)
        return f"""问题
{asked}

直接答案
{direct_answer}

关键依据
{evidence}

说明：千帆接口不可用时，系统已从分段文档中检索与问题最相关的原文，请人工核对后使用。"""
    if feature == "任务拆解":
        return f"""项目任务拆解

项目依据
{excerpt}

任务 1：确认目标与范围
工作说明：梳理项目目标、交付范围和限制条件，形成范围确认清单。
负责人建议：项目负责人｜优先级：高｜时间安排：第1个工作日
前置依赖：无
验收标准：目标、范围和限制条件均有明确记录，并经相关人员确认。
主要风险：目标表述不清可能导致后续任务方向偏差。

任务 2：整理资料与需求
工作说明：汇总现有资料，提取功能需求、业务规则和待确认事项。
负责人建议：产品人员｜优先级：高｜时间安排：第1—2个工作日
前置依赖：任务1完成
验收标准：需求清单内容完整，未确认信息已统一标记“待确认”。
主要风险：资料缺失可能造成需求遗漏。

任务 3：执行核心工作
工作说明：依据已确认的目标与需求完成核心任务，并记录执行过程。
负责人建议：执行人员｜优先级：高｜时间安排：第2—4个工作日
前置依赖：任务2完成
验收标准：核心工作按需求完成，过程和结果均可检查。
主要风险：资源不足或外部接口异常可能导致进度延期。

任务 4：检查并修订结果
工作说明：检查成果的完整性、准确性和可用性，并修复发现的问题。
负责人建议：测试人员｜优先级：中｜时间安排：第4个工作日
前置依赖：任务3完成
验收标准：主要问题均已修复，关键验收项检查通过。
主要风险：检查范围不足可能遗漏缺陷。

任务 5：交付与复盘
工作说明：整理最终成果，完成交付并记录问题、结论和后续改进事项。
负责人建议：项目负责人｜优先级：中｜时间安排：第5个工作日
前置依赖：任务4完成
验收标准：成果完成归档和交付，复盘记录包含问题、原因及改进建议。
主要风险：交付材料不完整可能影响验收。

风险提示
以上负责人、优先级和时间安排属于系统建议值；若原始资料未明确指定，应由项目负责人审核确认后执行。"""
    return f"""本周工作周报

一、本周工作概述
本周围绕既定目标推进相关工作，完成了资料整理、任务执行与阶段性结果确认。

二、本周完成事项
1. 梳理并归纳现有资料；
2. 推进重点任务并跟踪进度；
3. 完成阶段性成果检查与沟通。

三、工作成果
{excerpt}

四、存在问题
部分事项仍需进一步确认，后续应持续关注时间安排和协作效率。

五、下周计划
1. 完善未完成事项；
2. 根据反馈优化现有成果；
3. 明确下一阶段任务、负责人和截止时间。"""


def get_qianfan_config():
    """Read credentials without ever embedding them in source code."""
    api_key = os.getenv("QIANFAN_API_KEY", "").strip()
    model = os.getenv("QIANFAN_MODEL", "").strip()
    app_id = os.getenv("QIANFAN_APP_ID", "").strip()
    try:
        api_key = api_key or str(st.secrets.get("QIANFAN_API_KEY", "")).strip()
        model = model or str(st.secrets.get("QIANFAN_MODEL", "")).strip()
        app_id = app_id or str(st.secrets.get("QIANFAN_APP_ID", "")).strip()
    except (FileNotFoundError, KeyError):
        pass
    return api_key, model or "ernie-4.5-turbo-20260402", app_id


def build_ai_prompt(feature, text, question=""):
    instructions = {
        "会议纪要": """根据资料生成规范会议纪要，必须包含会议摘要、关键结论、待办事项、
负责人建议、优先级、截止时间和风险提示。资料中未出现的信息不要编造，并明确标记“待确认”。""",
        "文档问答": f"""仅根据资料回答问题：“{question.strip() or '这份资料的主要内容是什么？'}”
先给出直接答案，再列出关键依据并引用相关原文。引用时尽量保留资料中的文件名、页码、
幻灯片编号或段落编号；资料中没有答案时明确说明无法确定。""",
        "任务拆解": """根据资料将目标拆成可以执行的任务。每项任务包含任务名称、工作说明、
负责人建议、优先级、时间安排、前置依赖和验收标准，最后补充主要风险。""",
        "周报生成": """根据资料生成专业中文周报，包含本周工作概述、已完成事项、工作成果、
存在问题、下周计划和需要协调的事项。不要编造不存在的工作成果。""",
    }
    # Long documents are segmented and retrieved by the user's question. This
    # keeps relevant middle sections available instead of blindly deleting them.
    text = select_long_document_context(
        text,
        question if feature == "文档问答" else instructions[feature],
        30000,
    )
    return f"""安全要求：
不得输出完整身份证号、银行卡号、手机号、密码或详细住址；不得协助越权访问、泄露机密、歧视、诽谤、破坏性删除或隐瞒安全事件。遇到高风险请求时应明确拒绝，并给出合规替代方案。

任务要求：
{instructions[feature]}

用户资料：
{text}

请直接输出结构清晰、可编辑和可导出的中文结果，不要解释生成过程。"""


def generate_result(feature, text, question=""):
    """Use Qianfan when configured, with a safe local fallback for the demo."""
    safe_text, safe_question, blocked_response = prepare_safe_input(text, question)
    if blocked_response:
        st.session_state.ai_error = ""
        st.session_state.ai_provider = "本地安全策略"
        return blocked_response

    api_key, model, app_id = get_qianfan_config()
    st.session_state.ai_error = ""
    if not api_key:
        st.session_state.ai_provider = "本地演示"
        log_event("local_fallback_used", feature, True, detail="qianfan_not_configured")
        return generate_local_result(feature, safe_text, safe_question)

    try:
        client = OpenAI(
            api_key=api_key,
            base_url="https://qianfan.baidubce.com/v2",
            default_headers={"appid": app_id} if app_id else None,
            timeout=60.0,
            max_retries=2,
        )
        response = client.chat.completions.create(
            model=model,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "你是百度办公 AI 助手。你的回答应准确、简洁、结构化，"
                        "严格依据用户资料，不虚构姓名、日期、数据或结论。"
                    ),
                },
                {
                    "role": "user",
                    "content": build_ai_prompt(feature, safe_text, safe_question),
                },
            ],
            temperature=0.3,
            max_tokens=3000,
        )
        content = response.choices[0].message.content
        if not content:
            raise ValueError("模型返回了空内容")
        st.session_state.ai_provider = f"文心千帆 · {model}"
        log_event("qianfan_succeeded", feature, True)
        return redact_sensitive_text(content.strip())
    except Exception as exc:
        st.session_state.ai_provider = "本地演示（千帆调用失败）"
        st.session_state.ai_error = str(exc)
        log_event("qianfan_failed", feature, False, detail=type(exc).__name__)
        log_event("local_fallback_used", feature, True, detail=type(exc).__name__)
        return generate_local_result(feature, safe_text, safe_question)


def recommend_feature(text):
    """Recommend one office workflow, using Qianfan when available."""
    safe_text, _, blocked_response = prepare_safe_input(text)
    if blocked_response:
        return "文档问答", "资料包含高风险请求，建议先查看安全提示并按合规流程处理。", "本地安全策略"
    text = safe_text
    api_key, model, app_id = get_qianfan_config()
    if api_key:
        try:
            client = OpenAI(
                api_key=api_key,
                base_url="https://qianfan.baidubce.com/v2",
                default_headers={"appid": app_id} if app_id else None,
                timeout=30.0,
                max_retries=1,
            )
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {
                        "role": "system",
                        "content": "你是办公任务分类器，只能从会议纪要、文档问答、任务拆解、周报生成中选择一项。",
                    },
                    {
                        "role": "user",
                        "content": f"判断下面资料最适合哪项功能，并用“功能名称|一句理由”格式回答：\n{text[:8000]}",
                    },
                ],
                temperature=0,
                max_tokens=100,
            )
            raw = (response.choices[0].message.content or "").strip()
            for feature in FEATURES:
                if feature in raw:
                    reason = raw.split("|", 1)[1].strip() if "|" in raw else "千帆根据资料语义推荐此功能。"
                    return feature, reason, "文心千帆"
        except Exception:
            pass

    lowered = text.lower()
    rules = [
        ("会议纪要", ["会议", "参会", "讨论", "议程", "发言", "纪要"]),
        ("周报生成", ["本周", "下周", "周报", "工作进展", "完成情况"]),
        ("任务拆解", ["项目", "目标", "计划", "负责人", "截止", "任务"]),
        ("文档问答", ["报告", "文档", "白皮书", "合同", "制度", "说明"]),
    ]
    scores = {name: sum(lowered.count(word) for word in words) for name, words in rules}
    feature = max(scores, key=scores.get)
    if scores[feature] == 0:
        feature = "文档问答"
    return feature, f"资料中与“{feature}”相关的内容特征最明显。", "本地智能规则"


def chat_answer(messages, source_text=""):
    latest_question = ""
    if messages:
        latest_question = str(messages[-1].get("content", ""))
    safe_source, safe_question, blocked_response = prepare_safe_input(source_text, latest_question)
    if blocked_response:
        return blocked_response

    safe_messages = []
    for message in messages[-12:]:
        safe_messages.append({
            "role": message.get("role", "user"),
            "content": redact_sensitive_text(str(message.get("content", ""))),
        })
    if safe_messages and safe_messages[-1]["role"] == "user":
        safe_messages[-1]["content"] = safe_question

    api_key, model, app_id = get_qianfan_config()
    def local_reply():
        st.session_state.ai_provider = "本地咨询降级"
        if safe_source:
            passages = retrieve_relevant_passages(safe_source, safe_question, 3)
            evidence = "\n".join(f"- {passage}" for passage in passages)
            return f"""当前千帆接口暂不可用，已切换到本地资料检索。

针对你的问题“{safe_question}”，相关资料如下：
{evidence}

你可以继续围绕这些内容追问；正式使用前请结合原文核对。"""
        return f"""当前千帆接口暂不可用，已切换到本地咨询模式。

针对“{safe_question}”，建议先明确目标、范围、负责人、时间节点和验收标准，再拆分执行步骤并持续跟踪风险。你可以继续补充背景，我会结合上下文给出更具体的建议。"""

    if not api_key:
        return local_reply()
    context = safe_source[:16000] if safe_source else "用户尚未上传资料，可以回答通用办公问题。"
    try:
        client = OpenAI(
            api_key=api_key,
            base_url="https://qianfan.baidubce.com/v2",
            default_headers={"appid": app_id} if app_id else None,
            timeout=45.0,
            max_retries=1,
        )
        response = client.chat.completions.create(
            model=model,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "你是百度办公 AI 助手，支持多轮自然语言咨询。问题涉及资料时，只依据资料回答；"
                        "不得泄露敏感信息或协助越权、破坏、歧视、诽谤、机密外泄和隐瞒安全事件；"
                        f"资料如下：\n{context}"
                    ),
                }
            ] + safe_messages,
            temperature=0.3,
            max_tokens=1500,
        )
        st.session_state.ai_provider = f"文心千帆 · {model}"
        return redact_sensitive_text((response.choices[0].message.content or "未获得有效回复。 ").strip())
    except Exception as exc:
        log_event("chat_fallback_used", "AI 咨询", True, detail=type(exc).__name__)
        return local_reply()


def build_local_automation_tasks(goal, fallback_reason=""):
    """Create a complete, editable task plan when Function Calling is unavailable."""
    reason_note = f"；接口降级原因：{fallback_reason}" if fallback_reason else ""
    return [
        {
            "title": "确认目标、范围与验收口径",
            "owner": "项目负责人",
            "priority": "高",
            "deadline": "第 1 个工作日",
            "action": "确认目标、交付范围、资源约束和最终验收口径" + reason_note,
            "dependency": "无",
            "acceptance": "目标、范围、负责人和验收口径均已形成书面记录",
            "risk": "目标不清晰会导致后续任务返工",
            "done": False,
        },
        {
            "title": "拆分并分配核心任务",
            "owner": "项目经理",
            "priority": "高",
            "deadline": "第 2 个工作日",
            "action": f"根据目标拆分可执行任务并明确分工：{source_excerpt(goal, 120)}",
            "dependency": "目标、范围与验收口径已确认",
            "acceptance": "每项任务均包含负责人、优先级、截止时间和执行动作",
            "risk": "任务拆分过粗或责任边界不清",
            "done": False,
        },
        {
            "title": "执行任务并同步状态",
            "owner": "执行成员",
            "priority": "高",
            "deadline": "第 3—5 个工作日",
            "action": "按照任务清单执行工作，及时更新进度并记录阻塞问题",
            "dependency": "任务清单已确认并完成分配",
            "acceptance": "任务状态和成果记录完整，阻塞问题已明确责任人",
            "risk": "资源不足或依赖延迟可能影响进度",
            "done": False,
        },
        {
            "title": "验证结果并完成交付复盘",
            "owner": "审核成员",
            "priority": "中",
            "deadline": "第 6—7 个工作日",
            "action": "按照验收口径检查结果，修正问题并形成交付与复盘记录",
            "dependency": "核心任务已执行完成",
            "acceptance": "验收项全部有结论，遗留问题已记录后续计划",
            "risk": "验证范围不足可能遗漏问题",
            "done": False,
        },
    ]


def build_automation_tasks(goal):
    safe_goal, _, blocked_response = prepare_safe_input(goal)
    if blocked_response:
        st.session_state.automation_provider = "本地安全策略"
        return [{
            "title": "安全检查未通过",
            "owner": "申请人/系统管理员",
            "priority": "高",
            "deadline": "执行前",
            "action": blocked_response,
            "dependency": "完成合规审批和权限确认",
            "acceptance": "请求已改为合法、可授权且可审计的操作",
            "risk": "未经授权执行可能造成数据、安全或法律风险",
            "done": False,
        }]
    goal = safe_goal
    api_key, model, app_id = get_qianfan_config()
    if not api_key:
        st.session_state.automation_provider = "本地模拟任务（未配置千帆）"
        log_event("automation_fallback_used", "任务自动化", True, detail="qianfan_not_configured")
        return build_local_automation_tasks(goal, "未配置千帆接口")
    tools = [{
        "type": "function",
        "function": {
            "name": "build_task_plan",
            "description": "创建可执行、可跟踪的办公任务计划",
            "parameters": {
                "type": "object",
                "properties": {
                    "tasks": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "title": {"type": "string"},
                                "owner": {"type": "string"},
                                "priority": {"type": "string", "enum": ["高", "中", "低"]},
                                "deadline": {"type": "string"},
                                "action": {"type": "string"},
                                "dependency": {"type": "string"},
                                "acceptance": {"type": "string"},
                                "risk": {"type": "string"},
                            },
                            "required": ["title", "owner", "priority", "deadline", "action", "dependency", "acceptance", "risk"],
                        },
                    }
                },
                "required": ["tasks"],
            },
        },
    }]
    try:
        client = OpenAI(
            api_key=api_key,
            base_url="https://qianfan.baidubce.com/v2",
            default_headers={"appid": app_id} if app_id else None,
            timeout=45.0,
            max_retries=1,
        )
        response = client.chat.completions.create(
            model=model,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "你是办公任务自动化助手，请调用工具生成具体、可执行的任务。"
                        "资料未明确的信息应写待确认，不得虚构负责人和日期。"
                    ),
                },
                {"role": "user", "content": goal[:12000]},
            ],
            tools=tools,
            tool_choice={"type": "function", "function": {"name": "build_task_plan"}},
            temperature=0.2,
        )
        calls = response.choices[0].message.tool_calls or []
        if not calls:
            raise ValueError("模型未返回结构化任务")
        tasks = json.loads(calls[0].function.arguments).get("tasks", [])
        if not tasks:
            raise ValueError("模型返回的任务清单为空")
        for task in tasks:
            task.setdefault("owner", "待确认")
            task.setdefault("priority", "中")
            task.setdefault("deadline", "待确认")
            task.setdefault("action", "待确认")
            task.setdefault("dependency", "待确认")
            task.setdefault("acceptance", "待确认")
            task.setdefault("risk", "待确认")
            task["done"] = False
        st.session_state.automation_provider = "文心千帆 Function Calling"
        return tasks
    except Exception as exc:
        st.session_state.automation_provider = "本地模拟任务（千帆调用失败）"
        log_event("automation_failed", "任务自动化", False, detail=type(exc).__name__)
        log_event("automation_fallback_used", "任务自动化", True, detail=type(exc).__name__)
        return build_local_automation_tasks(goal, type(exc).__name__)


def add_history():
    item = {
        "feature": st.session_state.selected_feature,
        "title": f"{st.session_state.selected_feature}：{(st.session_state.source_names or ['粘贴文本'])[0]}",
        "source": "、".join(st.session_state.source_names) or "粘贴文本",
        "time": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "result": st.session_state.result,
    }
    with sqlite3.connect(DB_PATH) as connection:
        cursor = connection.execute(
            """INSERT INTO history_records (feature, title, source, created_at, result)
               VALUES (?, ?, ?, ?, ?)""",
            (item["feature"], item["title"], item["source"], item["time"], item["result"]),
        )
    item["id"] = cursor.lastrowid
    st.session_state.history.insert(0, item)


def add_markdown_runs(paragraph, text):
    """Add basic Markdown emphasis as real Word run formatting."""
    text = text.strip()
    cursor = 0
    for match in re.finditer(r"\*\*(.+?)\*\*|`(.+?)`", text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor:match.start()])
        content = match.group(1) if match.group(1) is not None else match.group(2)
        run = paragraph.add_run(content)
        if match.group(1) is not None:
            run.bold = True
        else:
            run.font.name = "Consolas"
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def markdown_table_cells(line):
    """Return Markdown table cells, excluding the outer pipe characters."""
    stripped = line.strip().strip("|")
    return [re.sub(r"\s+", " ", cell).strip() for cell in stripped.split("|")]


def is_markdown_table_separator(line):
    cells = markdown_table_cells(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


def style_docx_table(table):
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "微软雅黑")
                    run.font.size = Pt(9.5)
            if row_index == 0:
                shading = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
                if shading is None:
                    from docx.oxml import OxmlElement
                    shading = OxmlElement("w:shd")
                    cell._tc.get_or_add_tcPr().append(shading)
                shading.set(qn("w:fill"), "E8EEF5")
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(31, 77, 120)


def make_docx(text):
    buffer = BytesIO()
    document = Document()
    document.add_heading(f"百度办公 AI 助手｜{st.session_state.selected_feature}", 0)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    lines = normalize_generated_text(text).splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            index += 1
            continue

        # A Markdown table starts with a row followed by a separator row.
        if (
            "|" in stripped
            and index + 1 < len(lines)
            and is_markdown_table_separator(lines[index + 1])
        ):
            headers = markdown_table_cells(stripped)
            rows = []
            index += 2
            while index < len(lines) and "|" in lines[index] and lines[index].strip():
                row = markdown_table_cells(lines[index])
                if len(row) < len(headers):
                    row += [""] * (len(headers) - len(row))
                rows.append(row[: len(headers)])
                index += 1
            table = document.add_table(rows=1, cols=len(headers))
            for column, value in enumerate(headers):
                add_markdown_runs(table.rows[0].cells[column].paragraphs[0], value)
            for values in rows:
                cells = table.add_row().cells
                for column, value in enumerate(values):
                    add_markdown_runs(cells[column].paragraphs[0], value)
            style_docx_table(table)
            document.add_paragraph()
            continue

        heading = re.match(r"^(#{1,6})\s*(.+)$", stripped)
        if heading:
            title = re.sub(r"^\*\*|\*\*$", "", heading.group(2).strip())
            document.add_heading(title, level=min(len(heading.group(1)), 3))
            index += 1
            continue

        bullet = re.match(r"^[-*+]\s+(.+)$", stripped)
        numbered = re.match(r"^\d+[.)、]\s*(.+)$", stripped)
        if bullet or numbered:
            paragraph = document.add_paragraph(style="List Bullet" if bullet else "List Number")
            add_markdown_runs(paragraph, (bullet or numbered).group(1))
            index += 1
            continue

        paragraph = document.add_paragraph()
        add_markdown_runs(paragraph, stripped)
        index += 1

    document.save(buffer)
    return buffer.getvalue()


def make_pdf(text):
    buffer = BytesIO()
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    style = ParagraphStyle(
        "Chinese",
        fontName="STSong-Light",
        fontSize=11,
        leading=18,
        alignment=TA_LEFT,
        spaceAfter=4,
    )
    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=20 * mm,
        leftMargin=20 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )
    story = [
        Paragraph(escape(f"百度办公 AI 助手｜{st.session_state.selected_feature}"), style),
        Spacer(1, 8),
    ]
    for line in text.splitlines():
        story.append(Paragraph(escape(line) or "&nbsp;", style))
    document.build(story)
    return buffer.getvalue()


init_database()
init_state()
load_persistent_content()
if not st.session_state.visit_logged:
    log_event("app_visit")
    st.session_state.visit_logged = True

st.markdown(
    """
    <style>
    :root { --primary:#5d6df4; --soft:#eef2ff; --text:#1d2433; --muted:#8b93a5; --border:#e3e7f0; }
    .stApp { background:linear-gradient(135deg,#f9fbff 0%,#eef3ff 62%,#f8faff 100%); color:var(--text); }
    [data-testid="stHeader"] { background:transparent; }
    [data-testid="stSidebar"] { background:rgba(255,255,255,.97); border-right:1px solid var(--border); }
    [data-testid="stSidebar"] > div:first-child { padding-top:1.25rem; }
    .block-container { max-width:1160px; padding-top:1.5rem; padding-bottom:1.5rem; }
    .brand { display:flex;align-items:center;gap:10px;padding:4px 2px 18px;font-size:18px;font-weight:750; }
    .brand-logo { display:grid;place-items:center;width:36px;height:36px;border-radius:12px;color:white;
      background:linear-gradient(135deg,#4968ff,#7656e8);box-shadow:0 8px 20px rgba(82,99,238,.25); }
    .hero { margin:20px 0 16px; }.hero h1 { margin:0 0 8px;font-size:32px; }.hero h1 span { color:var(--primary); }
    .hero p { color:var(--muted);font-size:14px; }
    .feature-card { min-height:145px;padding:20px;border:1px solid var(--border);border-radius:16px;background:white; }
    .feature-title { font-size:18px;font-weight:750;margin-bottom:8px; }.feature-icon { font-size:25px;margin-right:8px; }
    .feature-card p { color:var(--muted);font-size:14px;line-height:1.65;margin:5px 0; }
    .tags { color:#69738b;font-size:12px;background:#f5f6fa;border-radius:8px;padding:5px 8px;display:inline-block; }
    .upload-zone { padding:30px;border:1.5px dashed #cbd2e2;border-radius:18px;text-align:center;background:rgba(255,255,255,.45); }
    .upload-icon { font-size:42px; }.upload-zone h3 { margin:8px 0; }.upload-zone p { color:var(--muted); }
    .result-box { padding:24px;border:1px solid var(--border);border-radius:16px;background:white;white-space:pre-wrap;
      line-height:1.85;min-height:300px; }
    .status-ok { display:inline-grid;place-items:center;width:44px;height:44px;border-radius:50%;background:#72d86e;
      color:white;font-size:28px;margin-right:10px; }
    .history-head { display:grid;grid-template-columns:2.4fr 1.4fr 1fr .5fr;gap:12px;align-items:center; }
    .history-head { padding:14px;font-weight:700;background:#f2f4fb;border-radius:12px 12px 0 0; }
    .footer { margin-top:28px;text-align:center;color:#a1a8b7;font-size:12px;line-height:2; }
    .stButton>button,.stDownloadButton>button { border-radius:10px;font-weight:650;border:1px solid #dce1ed; }
    .stButton>button[kind="primary"],.stDownloadButton>button[kind="primary"] {
      border:0;background:linear-gradient(135deg,#6071f5,#6d5bea);color:white; }
    .stTextArea textarea,.stTextInput input { border-radius:12px;border-color:var(--border); }
    div[data-testid="stSidebarNav"] { display:none; }
    </style>
    """,
    unsafe_allow_html=True,
)


with st.sidebar:
    st.markdown(
        '<div class="brand"><div class="brand-logo">AI</div><span>百度办公 AI 助手</span></div>',
        unsafe_allow_html=True,
    )
    configured_key, configured_model, _ = get_qianfan_config()
    if configured_key:
        st.success(f"千帆凭证已配置\n\n{configured_model}")
    else:
        st.info("当前为本地演示模式\n\n配置 API Key 后自动启用千帆")
    nav_items = ["首页", "上传资料", "AI 咨询", "会议纪要", "文档问答", "任务拆解", "任务自动化", "周报生成"]
    nav_icons = {
        "首页": "🏠",
        "上传资料": "📤",
        "AI 咨询": "🤖",
        "会议纪要": "📝",
        "文档问答": "💬",
        "任务拆解": "🧩",
        "任务自动化": "⚙️",
        "周报生成": "📊",
    }
    current_page = st.session_state.page
    if current_page == "上传资料" and st.session_state.workflow_feature_active:
        # 由内容生成功能进入上传步骤时，保持原功能高亮，避免用户
        # 误以为已经离开会议纪要/文档问答/任务拆解/周报生成流程。
        current_nav = st.session_state.selected_feature
    elif current_page == "选择功能":
        current_nav = st.session_state.selected_feature
    elif current_page in nav_items:
        current_nav = current_page
    elif current_page in {"生成结果", "编辑确认", "导出保存"}:
        current_nav = st.session_state.selected_feature
    else:
        # “选择功能”和辅助页面不属于主导航，不能继续误选“首页”。
        current_nav = None
    # 页面跳转后主动同步导航显示。导航点击由回调在主体渲染前处理，
    # 不再在页面执行中途 st.rerun，从而保证“下一步”等按钮事件不丢失。
    st.session_state["nav_radio_v2"] = current_nav
    st.radio(
        "功能导航",
        nav_items,
        index=nav_items.index(current_nav) if current_nav in nav_items else None,
        format_func=lambda x: f"{nav_icons[x]}  {x}",
        label_visibility="collapsed",
        key="nav_radio_v2",
        on_change=handle_nav_change,
    )
    st.divider()
    for label, target in [
        ("🕘  历史记录", "历史记录"),
        ("📁  我的文件", "我的文件"),
        ("📈  数据指标", "数据指标"),
        ("❓  帮助与反馈", "帮助与反馈"),
    ]:
        if st.button(
            label,
            key=f"side_{target}",
            type="primary" if current_page == target else "secondary",
            use_container_width=True,
        ):
            go(target)
            st.rerun()


st.caption("智能办公 · 高效协作")


def footer():
    st.markdown(
        '<div class="footer">百度办公 AI 助手　|　自然语言咨询　|　智能推荐　|　任务自动化　|　办公内容生成<br>'
        '© 2026 百度办公 AI 助手，本界面仅用于实习项目原型展示</div>',
        unsafe_allow_html=True,
    )


page = st.session_state.page

if page == "首页":
    st.markdown(
        '<div class="hero"><h1>你好</h1>'
        '<p>上传资料或输入问题，让 AI 帮你更高效地完成工作</p></div>',
        unsafe_allow_html=True,
    )
    home_text = st.text_area(
        "快速输入",
        placeholder="粘贴会议记录、项目资料或本周工作内容，快速开始办公处理……",
        label_visibility="collapsed",
        height=130,
    )
    left, middle, space, right = st.columns([1.4, 1.4, 4, 1.2])
    with left:
        if st.button("📤 上传资料", use_container_width=True):
            go("上传资料")
            st.rerun()
    with middle:
        if st.button("📋 使用文本", use_container_width=True):
            if home_text.strip():
                st.session_state.source_text = home_text.strip()
                st.session_state.source_names = ["首页粘贴文本"]
                go("选择功能")
                st.rerun()
            st.warning("请先输入文本。")
    with right:
        if st.button("开始 ➜", type="primary", use_container_width=True):
            if home_text.strip():
                st.session_state.source_text = home_text.strip()
                st.session_state.source_names = ["首页粘贴文本"]
                go("选择功能")
                st.rerun()
            go("上传资料")
            st.rerun()
    st.markdown("### 快捷功能")
    card_cols = st.columns(2, gap="large")
    for index, (name, info) in enumerate(FEATURES.items()):
        with card_cols[index % 2]:
            st.markdown(
                f'<div class="feature-card"><div class="feature-title"><span class="feature-icon">{info["icon"]}</span>{name}</div>'
                f'<p>{info["desc"]}</p><span class="tags">{info["tags"]}</span></div>',
                unsafe_allow_html=True,
            )
            if st.button(f"进入{name}", key=f"home_{name}", use_container_width=True):
                st.session_state.selected_feature = name
                st.session_state.workflow_feature_active = True
                go("上传资料")
                st.rerun()
            st.write("")

elif page == "上传资料":
    st.markdown('<div class="hero"><h1>上传资料</h1><p>支持上传文档或者直接粘贴文本，快速开始 AI 办公处理。</p></div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="upload-zone"><div class="upload-icon">⇧</div><h3>拖拽文件到这里，或点击下方选择文件</h3>'
        '<p>支持 PDF、Word、PowerPoint、TXT 和 Markdown；单个文件不超过 20MB，单次最多 3 个</p></div>',
        unsafe_allow_html=True,
    )
    st.warning("请勿上传身份证号、银行卡号、密码、合同机密或其他敏感信息。资料会发送至千帆模型接口处理；MVP 日志不会保存文档正文。")
    uploads = st.file_uploader(
        "选择资料",
        type=["pdf", "docx", "pptx", "txt", "md"],
        accept_multiple_files=True,
        label_visibility="collapsed",
    )
    pasted = st.text_area("或粘贴文本", placeholder="也可以直接粘贴会议记录、项目目标或本周工作内容……", height=150)
    if uploads:
        for file in uploads:
            size = getattr(file, "size", len(file.getvalue()))
            status = "✓ 待解析" if size <= MAX_UPLOAD_BYTES else "✕ 超过 20MB"
            st.caption(f"{status}　{file.name}　({size / 1024 / 1024:.2f} MB)")
        if len(uploads) > MAX_UPLOAD_FILES:
            st.error(f"单次最多上传 {MAX_UPLOAD_FILES} 个文件，当前已选择 {len(uploads)} 个。")
    back, space, next_col = st.columns([1, 5, 1.5])
    with back:
        if st.button("返回首页", use_container_width=True):
            go("首页")
            st.rerun()
    with next_col:
        if st.button("下一步：选择功能 ➜", type="primary", use_container_width=True):
            pieces, names = [], []
            if len(uploads or []) > MAX_UPLOAD_FILES:
                st.error(f"单次最多上传 {MAX_UPLOAD_FILES} 个文件，请删除多余文件后重试。")
                st.stop()
            oversized = [
                file.name for file in uploads or []
                if getattr(file, "size", len(file.getvalue())) > MAX_UPLOAD_BYTES
            ]
            if oversized:
                st.error("以下文件超过 20MB，无法处理：" + "、".join(oversized))
                st.stop()
            try:
                for file in uploads or []:
                    text = extract_file(file)
                    if text.strip():
                        pieces.append(f"【{file.name}】\n{text}")
                    else:
                        st.warning(f"{file.name} 未提取到可读文字；如果是扫描型 PDF，请上传可复制文本版本。")
                    names.append(file.name)
            except Exception as exc:
                log_event("upload_failed", success=False, detail=type(exc).__name__)
                st.error(f"文件解析失败：{exc}")
                st.stop()
            if pasted.strip():
                pieces.append(f"【粘贴文本】\n{pasted.strip()}")
                names.append("粘贴文本")
            if not pieces:
                st.warning("请至少上传一个可读取的文件或粘贴一段文本。")
            else:
                st.session_state.source_text = "\n\n".join(pieces)
                st.session_state.source_names = names
                st.session_state.workflow_feature_active = True
                log_event("upload_success", success=True, detail=f"files={len(names)}")
                recommended, reason, provider = recommend_feature(st.session_state.source_text)
                st.session_state.recommended_feature = recommended
                st.session_state.recommendation_reason = f"{reason}（推荐来源：{provider}）"
                go("选择功能")
                st.rerun()

elif page == "选择功能":
    st.markdown('<div class="hero"><h1>选择处理功能</h1><p>根据上传的资料内容，选择一项需要处理的办公任务。</p></div>', unsafe_allow_html=True)
    if not st.session_state.recommended_feature:
        recommended, reason, provider = recommend_feature(st.session_state.source_text)
        st.session_state.recommended_feature = recommended
        st.session_state.recommendation_reason = f"{reason}（推荐来源：{provider}）"
    recommendation = st.session_state.recommended_feature
    st.info(f"✨ 智能推荐：**{recommendation}**\n\n{st.session_state.recommendation_reason}")
    if st.button(f"采用推荐：{recommendation}", type="primary"):
        st.session_state.selected_feature = recommendation
        log_event("recommendation_adopted", recommendation, True)
        st.rerun()
    cols = st.columns(2, gap="large")
    for index, (name, info) in enumerate(FEATURES.items()):
        with cols[index % 2]:
            selected = st.session_state.selected_feature == name
            st.markdown(
                f'<div class="feature-card"><div class="feature-title"><span class="feature-icon">{info["icon"]}</span>{name}</div>'
                f'<p>{info["desc"]}</p><span class="tags">{info["tags"]}</span></div>',
                unsafe_allow_html=True,
            )
            if st.button("✓ 已选择" if selected else f"选择{name}", key=f"choose_{name}", type="primary" if selected else "secondary", use_container_width=True):
                st.session_state.selected_feature = name
                log_event("feature_selected", name, True)
                st.rerun()
            st.write("")
    if st.session_state.selected_feature == "文档问答":
        st.session_state.question = st.text_input(
            "请输入你想围绕资料询问的问题",
            value=st.session_state.question,
            placeholder="例如：这份文档的主要结论是什么？",
        )
    back, space, start = st.columns([1.2, 5, 1.4])
    with back:
        if st.button("重新上传", use_container_width=True):
            go("上传资料")
            st.rerun()
    with start:
        if st.button("开始生成 ➜", type="primary", use_container_width=True):
            feature = st.session_state.selected_feature
            log_event("generation_started", feature)
            started = time.perf_counter()
            with st.spinner("AI 正在生成结果，请稍候……"):
                st.session_state.result = normalize_generated_text(
                    generate_result(feature, st.session_state.source_text, st.session_state.question)
                )
            st.session_state.result_confirmed = False
            duration_ms = int((time.perf_counter() - started) * 1000)
            st.session_state.last_generation_ms = duration_ms
            result_returned = bool(st.session_state.result.strip())
            log_event("result_displayed", feature, result_returned, duration_ms, st.session_state.ai_provider)
            log_event("generation_success" if result_returned else "generation_failed", feature, result_returned, duration_ms, st.session_state.ai_error)
            add_history()
            go("生成结果")
            st.rerun()

elif page == "生成结果":
    st.session_state.result = normalize_generated_text(st.session_state.result)
    left, right = st.columns([5, 2])
    with left:
        st.markdown(
            f'<div class="hero"><h1><span class="status-ok">✓</span>生成结果</h1>'
            f'<p>{st.session_state.selected_feature}　基于当前上传资料生成</p></div>',
            unsafe_allow_html=True,
        )
    with right:
        st.write("")
        c1, c2 = st.columns(2)
        with c1:
            if st.button("⟳ 重新生成", use_container_width=True):
                feature = st.session_state.selected_feature
                log_event("regenerate_clicked", feature, True)
                started = time.perf_counter()
                with st.spinner("AI 正在重新生成，请稍候……"):
                    st.session_state.result = normalize_generated_text(
                        generate_result(feature, st.session_state.source_text, st.session_state.question)
                    )
                st.session_state.result_confirmed = False
                duration_ms = int((time.perf_counter() - started) * 1000)
                st.session_state.last_generation_ms = duration_ms
                result_returned = bool(st.session_state.result.strip())
                log_event("result_displayed", feature, result_returned, duration_ms, st.session_state.ai_provider)
                log_event("generation_success" if result_returned else "generation_failed", feature, result_returned, duration_ms, st.session_state.ai_error)
                st.rerun()
        with c2:
            if st.button("⇩ 导出结果", type="primary", use_container_width=True):
                if st.session_state.result_confirmed:
                    go("导出保存")
                else:
                    st.session_state.confirmation_notice = True
                    go("编辑确认")
                st.rerun()
    if st.session_state.ai_error:
        with st.expander("千帆调用失败，已自动使用本地演示结果", expanded=True):
            st.error(st.session_state.ai_error)
            st.caption("请检查 API Key、模型名称、接入点权限、账户余额和网络连接。")
    else:
        st.caption(f"生成服务：{st.session_state.ai_provider}")
    st.warning("AI 生成内容仅供参考，请核对姓名、日期、数据和关键结论后再使用。")
    result_col, info_col = st.columns([2.2, 1], gap="large")
    with result_col:
        # Render model Markdown natively. Displaying Markdown tables inside a
        # pre-wrapped HTML block preserves their alignment spaces and can
        # create very large invisible vertical gaps when lines wrap.
        with st.container(border=True):
            st.markdown(st.session_state.result)
        edit_col, export_col = st.columns(2)
        with edit_col:
            if st.button("✎ 编辑确认", use_container_width=True):
                go("编辑确认")
                st.rerun()
        with export_col:
            if st.button("导出/保存", type="primary", use_container_width=True):
                if st.session_state.result_confirmed:
                    go("导出保存")
                else:
                    st.session_state.confirmation_notice = True
                    go("编辑确认")
                st.rerun()
    with info_col:
        st.markdown("### 参考资料")
        with st.container(border=True):
            for name in st.session_state.source_names:
                st.write(f"📄 {name}")
        st.markdown("### 本次操作")
        with st.container(border=True):
            st.write(f"◷ 生成时间：{datetime.now():%Y-%m-%d}")
            st.write(f"💬 处理类型：{st.session_state.selected_feature}")
            st.write(f"♙ 资料数量：{len(st.session_state.source_names)}")
            if st.session_state.last_generation_ms:
                st.write(f"⏱ 生成耗时：{st.session_state.last_generation_ms / 1000:.1f} 秒")

elif page == "编辑确认":
    st.markdown('<div class="hero"><h1>编辑确认</h1><p>请检查并修改以下内容，确认无误后保存或导出结果。</p></div>', unsafe_allow_html=True)
    if st.session_state.confirmation_notice:
        st.warning("导出前请人工核对内容。完成核对后，可在本页保存修改或直接导出。")
        st.session_state.confirmation_notice = False
    main, side = st.columns([2.2, 1], gap="large")
    with main:
        st.caption("来源资料：" + ("、".join(st.session_state.source_names) or "粘贴文本"))
        edited = st.text_area("AI 生成结果（可编辑）", value=st.session_state.result, height=430)
    with side:
        st.markdown("### 参考资料")
        with st.container(border=True):
            for name in st.session_state.source_names:
                st.write(f"📄 {name}")
        st.markdown("### 编辑提示")
        st.info("请重点核对姓名、日期、任务负责人和关键结论，避免 AI 内容出现遗漏。")
    cancel, space, regenerate, export, save = st.columns([1, 3, 1.4, 1.4, 1.4])
    with cancel:
        if st.button("取消", use_container_width=True):
            go("生成结果")
            st.rerun()
    with regenerate:
        if st.button("⟳ 重新生成", use_container_width=True):
            st.session_state.result = normalize_generated_text(
                generate_result(
                    st.session_state.selected_feature,
                    st.session_state.source_text,
                    st.session_state.question,
                )
            )
            st.session_state.result_confirmed = False
            st.rerun()
    with export:
        if st.button("⇩ 导出结果", type="primary", use_container_width=True):
            st.session_state.result = edited
            st.session_state.result_confirmed = True
            log_event("result_confirmed", st.session_state.selected_feature, True, detail="export_from_review")
            go("导出保存")
            st.rerun()
    with save:
        if st.button("✓ 保存修改", type="primary", use_container_width=True):
            changed = edited != st.session_state.result
            st.session_state.result = edited
            st.session_state.result_confirmed = True
            if st.session_state.history:
                st.session_state.history[0]["result"] = edited
                history_id = st.session_state.history[0].get("id")
                if history_id is not None:
                    with sqlite3.connect(DB_PATH) as connection:
                        connection.execute("UPDATE history_records SET result = ? WHERE id = ?", (edited, history_id))
            log_event("edit_saved", st.session_state.selected_feature, True, detail=f"changed={int(changed)}")
            log_event("result_confirmed", st.session_state.selected_feature, True, detail="saved_from_review")
            st.success("修改已保存。")

elif page == "导出保存":
    st.markdown('<div class="hero"><h1>导出/保存结果</h1><p>选择导出格式或保存位置。</p></div>', unsafe_allow_html=True)
    if not st.session_state.result_confirmed:
        st.warning("当前结果尚未人工确认，请先进入编辑确认页面完成核对。")
        if st.button("前往编辑确认", type="primary"):
            go("编辑确认")
            st.rerun()
        st.stop()
    format_name = st.radio(
        "导出格式",
        ["Word (.docx)", "PDF (.pdf)", "TXT (.txt)"],
        horizontal=True,
    )
    default_name = f"{st.session_state.selected_feature}结果"
    filename = st.text_input("文件名称", value=default_name)
    if format_name.startswith("Word"):
        file_data, suffix, mime = make_docx(st.session_state.result), "docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif format_name.startswith("PDF"):
        file_data, suffix, mime = make_pdf(st.session_state.result), "pdf", "application/pdf"
    else:
        file_data, suffix, mime = st.session_state.result.encode("utf-8-sig"), "txt", "text/plain"
    st.markdown("### 内容预览")
    with st.container(border=True):
        st.markdown(st.session_state.result[:1000])
    cancel, space, save_col, download_col = st.columns([1, 4, 1.7, 1.5])
    with cancel:
        if st.button("取消", use_container_width=True):
            go("生成结果")
            st.rerun()
    with save_col:
        if st.button("▣ 保存到我的文件", use_container_width=True):
            item = persist_saved_file(
                {"name": f"{filename}.{suffix}", "data": file_data, "mime": mime, "time": datetime.now().strftime("%Y-%m-%d %H:%M")}
            )
            st.session_state.saved_files.append(item)
            log_event("result_saved", st.session_state.selected_feature, True, detail=suffix)
            st.success("已保存到“我的文件”。")
    with download_col:
        st.download_button(
            "⇩ 导出文件",
            data=file_data,
            file_name=f"{filename}.{suffix}",
            mime=mime,
            type="primary",
            use_container_width=True,
            on_click=log_event,
            args=("result_exported", st.session_state.selected_feature, True, None, suffix),
        )

elif page == "AI 咨询":
    st.markdown('<div class="hero"><h1>自然语言咨询</h1><p>支持多轮连续对话，也可以结合当前上传资料回答问题。</p></div>', unsafe_allow_html=True)
    status_col, action_col = st.columns([5, 1])
    with status_col:
        if st.session_state.source_text:
            st.success("已关联资料：" + ("、".join(st.session_state.source_names) or "粘贴文本"))
        else:
            st.info("当前为通用办公咨询。上传资料后，可围绕资料连续提问。")
    with action_col:
        if st.button("清空对话", use_container_width=True):
            st.session_state.chat_messages = [{"role": "assistant", "content": "对话已清空，请输入新的问题。"}]
            log_event("chat_cleared", success=True)
            st.rerun()
    for message in st.session_state.chat_messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
    prompt = st.chat_input("输入办公问题，或围绕已上传资料继续提问")
    if prompt:
        st.session_state.chat_messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)
        try:
            started = time.perf_counter()
            with st.chat_message("assistant"):
                with st.spinner("正在结合上下文思考……"):
                    reply = chat_answer(st.session_state.chat_messages, st.session_state.source_text)
                st.markdown(reply)
            duration_ms = int((time.perf_counter() - started) * 1000)
            st.session_state.chat_messages.append({"role": "assistant", "content": reply})
            log_event("chat_success", "AI 咨询", True, duration_ms)
        except Exception as exc:
            st.error(f"咨询失败：{exc}")
            log_event("chat_failed", "AI 咨询", False, detail=type(exc).__name__)

elif page == "任务自动化":
    st.markdown('<div class="hero"><h1>任务自动化</h1><p>通过 Function Calling 创建结构化任务，并持续跟踪完成状态。</p></div>', unsafe_allow_html=True)
    st.info(
        "本功能仅用于生成任务清单、辅助规划和进度跟踪，不连接项目管理、邮件、日历等外部办公平台，"
        "也不会替用户真实创建、分派或执行外部任务。所有任务须经人工确认后再实施。"
    )
    default_goal = source_excerpt(st.session_state.source_text, 1200) if st.session_state.source_text else ""
    automation_goal = st.text_area(
        "项目目标或工作要求",
        value=default_goal,
        placeholder="例如：两周内完成校园活动策划，包括方案、预算、宣传和现场分工。",
        height=150,
    )
    run_col, clear_col = st.columns([1.6, 1])
    with run_col:
        if st.button("⚙️ 自动生成任务", type="primary", use_container_width=True):
            if not automation_goal.strip():
                st.warning("请先输入项目目标。")
            else:
                try:
                    started = time.perf_counter()
                    with st.spinner("正在调用任务规划工具……"):
                        clear_automation_checkbox_state()
                        st.session_state.automation_tasks = build_automation_tasks(automation_goal)
                    log_event("automation_created", "任务自动化", True, int((time.perf_counter() - started) * 1000), f"tasks={len(st.session_state.automation_tasks)}")
                    st.rerun()
                except Exception as exc:
                    st.error(f"任务自动化失败：{exc}")
                    log_event("automation_failed", "任务自动化", False, detail=type(exc).__name__)
    with clear_col:
        if st.button("清空任务", use_container_width=True):
            clear_automation_checkbox_state()
            st.session_state.automation_tasks = []
            log_event("automation_cleared", "任务自动化", True)
            st.rerun()
    tasks = st.session_state.automation_tasks
    if tasks:
        if st.session_state.automation_provider.startswith("本地模拟任务"):
            st.warning(f"任务来源：{st.session_state.automation_provider}。请人工核对后使用。")
        else:
            st.caption(f"任务来源：{st.session_state.automation_provider}")
        completed = sum(bool(task.get("done")) for task in tasks)
        st.progress(completed / len(tasks), text=f"总体进度：{completed}/{len(tasks)} 项已完成")
        for index, task in enumerate(tasks):
            with st.container(border=True):
                done_col, info_col, priority_col = st.columns([0.5, 5, 1.2])
                with done_col:
                    st.checkbox(
                        "完成",
                        value=bool(task.get("done")),
                        key=f"automation_done_{index}",
                        label_visibility="collapsed",
                        on_change=update_automation_task_status,
                        args=(index,),
                    )
                with info_col:
                    st.markdown(f"**{task['title']}**")
                    st.caption(f"负责人：{task['owner']}　｜　截止：{task['deadline']}")
                    st.write(task["action"])
                    st.caption(
                        f"前置依赖：{task.get('dependency', '待确认')}\n\n"
                        f"验收标准：{task.get('acceptance', '待确认')}\n\n"
                        f"主要风险：{task.get('risk', '待确认')}"
                    )
                with priority_col:
                    priority_icon = {"高": "🔴", "中": "🟠", "低": "🟢"}.get(task["priority"], "⚪")
                    st.write(f"{priority_icon} {task['priority']}")
        task_text = "\n".join(
            f"{'✓' if task.get('done') else '□'} {task['title']}｜负责人：{task['owner']}｜优先级：{task['priority']}｜"
            f"截止：{task['deadline']}｜执行动作：{task['action']}｜前置依赖：{task.get('dependency', '待确认')}｜"
            f"验收标准：{task.get('acceptance', '待确认')}｜主要风险：{task.get('risk', '待确认')}"
            for task in tasks
        )
        st.download_button("导出任务清单", task_text.encode("utf-8-sig"), "自动化任务清单.txt", "text/plain", use_container_width=True)
    else:
        st.info("输入目标后，系统会生成包含负责人、优先级、截止时间、执行动作、前置依赖、验收标准和主要风险的任务清单。")

elif page == "数据指标":
    st.markdown('<div class="hero"><h1>数据指标</h1><p>基于匿名事件日志查看 MVP 核心流程表现，不保存上传文档正文。</p></div>', unsafe_allow_html=True)
    with sqlite3.connect(DB_PATH) as connection:
        rows = connection.execute("SELECT event_name, COUNT(*), SUM(CASE WHEN success=1 THEN 1 ELSE 0 END), AVG(duration_ms) FROM events GROUP BY event_name").fetchall()
        feedback_row = connection.execute("SELECT COUNT(*), AVG(rating), AVG(helpful), AVG(reuse_intent) FROM feedback").fetchone()
    event_stats = {
        name: {"count": count, "successes": successes or 0, "avg_ms": avg or 0}
        for name, count, successes, avg in rows
    }
    counts = {name: stats["count"] for name, stats in event_stats.items()}
    requests = counts.get("generation_started", 0)
    qianfan_successes = counts.get("qianfan_succeeded", 0)
    qianfan_failures = counts.get("qianfan_failed", 0)
    qianfan_attempts = qianfan_successes + qianfan_failures
    fallback_successes = event_stats.get("local_fallback_used", {}).get("successes", 0)
    displayed_results = event_stats.get("result_displayed", {}).get("successes", 0)
    real_model_rate = qianfan_successes / qianfan_attempts * 100 if qianfan_attempts else None
    fallback_rate = fallback_successes / requests * 100 if requests else None
    overall_return_rate = displayed_results / requests * 100 if requests else None
    api_failure_rate = qianfan_failures / qianfan_attempts * 100 if qianfan_attempts else None
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("访问次数", counts.get("app_visit", 0))
    c2.metric("生成请求", requests)
    c3.metric("总体结果返回率", f"{overall_return_rate:.1f}%" if overall_return_rate is not None else "暂无")
    avg_ms = event_stats.get("result_displayed", {}).get("avg_ms", 0)
    c4.metric("平均生成时间", f"{avg_ms / 1000:.1f} 秒" if avg_ms else "暂无")
    st.markdown("### AI调用与降级")
    ai1, ai2, ai3, ai4 = st.columns(4)
    ai1.metric("千帆真实成功率", f"{real_model_rate:.1f}%" if real_model_rate is not None else "暂无")
    ai2.metric("本地降级率", f"{fallback_rate:.1f}%" if fallback_rate is not None else "暂无")
    ai3.metric("API调用失败率", f"{api_failure_rate:.1f}%" if api_failure_rate is not None else "暂无")
    ai4.metric("降级成功次数", fallback_successes)
    st.caption("真实成功率仅统计千帆调用；总体结果返回率同时包含千帆成功和本地降级成功，二者不能混为同一指标。")
    st.markdown("### 转化与反馈")
    c5, c6, c7, c8 = st.columns(4)
    c5.metric("重新生成次数", counts.get("regenerate_clicked", 0))
    c6.metric("导出次数", counts.get("result_exported", 0))
    c7.metric("保存次数", counts.get("result_saved", 0))
    c8.metric("平均评分", f"{feedback_row[1]:.1f}/5" if feedback_row and feedback_row[1] else "暂无")
    st.markdown("### 事件明细")
    st.dataframe([{"事件": name, "次数": count, "平均耗时(ms)": round(avg or 0)} for name, count, _, avg in rows], use_container_width=True, hide_index=True)

elif page == "历史记录":
    st.markdown('<div class="hero"><h1>历史记录</h1><p>查看并复用曾经生成的办公内容。</p></div>', unsafe_allow_html=True)
    search_col, type_col, clear_col = st.columns([3, 1.3, 1])
    with search_col:
        search = st.text_input("搜索", placeholder="搜索记录内容", label_visibility="collapsed")
    with type_col:
        type_filter = st.selectbox("类型", ["全部类型"] + list(FEATURES), label_visibility="collapsed")
    with clear_col:
        if st.button("🗑 清空记录", use_container_width=True):
            clear_history_records()
            st.session_state.history = []
            st.rerun()
    records = [
        item for item in st.session_state.history
        if (not search or search.lower() in (item["title"] + item["result"]).lower())
        and (type_filter == "全部类型" or item["feature"] == type_filter)
    ]
    st.markdown('<div class="history-head"><div>内容</div><div>来源文件</div><div>生成时间</div><div>操作</div></div>', unsafe_allow_html=True)
    if not records:
        st.info("暂无符合条件的历史记录。完成一次内容生成后会自动保存在这里。")
    for index, item in enumerate(records):
        cols = st.columns([2.4, 1.4, 1, .8])
        cols[0].write(f"**{item['title']}**")
        cols[1].write(item["source"])
        cols[2].write(item["time"])
        with cols[3]:
            view_col, delete_col = st.columns(2)
            if view_col.button("查看", key=f"view_{index}"):
                st.session_state.result = normalize_generated_text(item["result"])
                st.session_state.selected_feature = item["feature"]
                go("生成结果")
                st.rerun()
            if delete_col.button("删除", key=f"delete_{index}"):
                delete_history_record(item.get("id"))
                st.session_state.history.remove(item)
                log_event("history_deleted", item["feature"], True)
                st.rerun()
        st.divider()

elif page == "我的文件":
    st.markdown('<div class="hero"><h1>我的文件</h1><p>下载已保存的生成结果。</p></div>', unsafe_allow_html=True)
    if not st.session_state.saved_files:
        st.info("暂无已保存文件。你可以在“导出/保存结果”页面将结果保存到这里。")
    for index, item in enumerate(reversed(st.session_state.saved_files)):
        cols = st.columns([3, 1.5, 1])
        cols[0].write(f"📄 **{item['name']}**")
        cols[1].write(item["time"])
        cols[2].download_button("下载", item["data"], item["name"], item["mime"], key=f"saved_{index}", use_container_width=True)

elif page == "帮助与反馈":
    st.markdown('<div class="hero"><h1>帮助与反馈</h1><p>了解系统流程，并提交真实使用反馈。</p></div>', unsafe_allow_html=True)
    with st.container(border=True):
        st.markdown(
            """
            **使用流程**

            1. 上传 PDF、Word、PowerPoint、TXT、Markdown，或直接粘贴文本；
            2. 选择会议纪要、文档问答、任务拆解或周报生成；
            3. 查看 AI 生成结果并人工编辑确认；
            4. 导出 Word、PDF、TXT，或保存到历史记录。

            **安全与使用说明**

            AI 生成内容仅供参考，正式使用前必须人工核对。请勿上传身份证号、银行卡号、密码、合同机密等敏感信息。上传内容会发送至千帆模型接口处理，MVP 指标日志只记录匿名 ID、文件类型、处理结果和耗时，不保存文档正文。

            """
        )
    st.markdown("### 提交反馈")
    with st.form("feedback_form", clear_on_submit=True):
        rating = st.slider("结果满意度", 1, 5, 4)
        helpful = st.radio("结果是否有帮助？", ["有帮助", "没有帮助"], horizontal=True)
        reuse_intent = st.radio("是否愿意再次使用？", ["愿意", "不愿意"], horizontal=True)
        comment = st.text_area("意见或问题", placeholder="请描述结果质量、页面流程或系统稳定性方面的问题……")
        if st.form_submit_button("提交反馈", type="primary"):
            save_feedback(rating, helpful == "有帮助", reuse_intent == "愿意", comment)
            st.success("感谢反馈，已匿名记录。")

footer()
